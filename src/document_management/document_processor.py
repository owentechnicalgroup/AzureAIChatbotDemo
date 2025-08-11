"""
Document processing module for RAG implementation.
Handles file upload, text extraction, and document chunking.
"""

import asyncio
import io
import os
import uuid
from pathlib import Path
from typing import List, Optional, Dict, Any, Union
from datetime import datetime, timezone

import structlog
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

from src.config.settings import Settings

logger = structlog.get_logger(__name__)


class DocumentProcessor:
    """
    Handles document processing including text extraction and chunking.
    
    Supports:
    - PDF files (using pypdf)
    - DOCX files (using python-docx)
    - TXT files (plain text)
    """
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    MAX_FILE_SIZE_BYTES = 100 * 1024 * 1024  # 100MB default
    
    def __init__(self, settings: Settings):
        """
        Initialize document processor with settings.
        
        Args:
            settings: Application settings containing chunk configuration
        """
        self.settings = settings
        self.logger = logger.bind(
            log_type="SYSTEM", 
            component="document_processor"
        )
        
        # Configure text splitter based on settings
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Update max file size from settings
        self.max_file_size_bytes = settings.max_file_size_mb * 1024 * 1024
        
        self.logger.info(
            "DocumentProcessor initialized",
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            max_file_size_mb=settings.max_file_size_mb
        )
    
    def validate_file(self, file_path: Union[str, Path], file_size: Optional[int] = None, allow_memory_only: bool = False) -> bool:
        """
        Validate file for processing.
        
        Args:
            file_path: Path to the file (used for extension detection)
            file_size: File size in bytes (optional, will be calculated if not provided)
            allow_memory_only: If True, skip file existence check for in-memory processing
            
        Returns:
            True if file is valid for processing
            
        Raises:
            ValueError: If file is invalid
        """
        file_path = Path(file_path)
        
        # Check if file exists (skip for in-memory processing)
        if not allow_memory_only and not file_path.exists():
            raise ValueError(f"File does not exist: {file_path}")
        
        # Check file extension
        if file_path.suffix.lower() not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: {file_path.suffix}. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
            )
        
        # Check file size
        if file_size is None and not allow_memory_only:
            file_size = file_path.stat().st_size
        elif file_size is None:
            # For in-memory processing, file_size should be provided
            raise ValueError("File size must be provided for in-memory processing")
        
        if file_size > self.max_file_size_bytes:
            raise ValueError(
                f"File too large: {file_size / (1024*1024):.1f}MB. "
                f"Maximum allowed: {self.max_file_size_bytes / (1024*1024)}MB"
            )
        
        self.logger.debug(
            "File validation passed",
            file_path=str(file_path),
            file_size_mb=file_size / (1024*1024),
            file_extension=file_path.suffix,
            memory_only=allow_memory_only
        )
        
        return True
    
    async def extract_text(self, file_path: Union[str, Path], file_content: Optional[bytes] = None) -> str:
        """
        Extract text from various file formats.
        
        Args:
            file_path: Path to the file (used for extension detection)
            file_content: File content as bytes (optional, will read from file_path if not provided)
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is unsupported
            Exception: If text extraction fails
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        try:
            # Read file content if not provided
            if file_content is None:
                with open(file_path, 'rb') as f:
                    file_content = f.read()
            
            # Create in-memory buffer
            buffer = io.BytesIO(file_content)
            
            if extension == '.pdf':
                text = await self._extract_pdf_text(buffer)
            elif extension == '.docx':
                text = await self._extract_docx_text(buffer)
            elif extension == '.txt':
                text = await self._extract_txt_text(file_content)
            else:
                raise ValueError(f"Unsupported file extension: {extension}")
            
            self.logger.info(
                "Text extraction completed",
                file_path=str(file_path),
                text_length=len(text),
                file_extension=extension
            )
            
            return text
            
        except Exception as e:
            self.logger.error(
                "Text extraction failed",
                file_path=str(file_path),
                error=str(e),
                error_type=type(e).__name__
            )
            raise
    
    async def _extract_pdf_text(self, buffer: io.BytesIO) -> str:
        """Extract text from PDF using pypdf."""
        try:
            reader = PdfReader(buffer)
            text_parts = []
            
            # Check if PDF is encrypted
            if reader.is_encrypted:
                self.logger.warning("PDF is encrypted/password-protected")
                raise ValueError("PDF is encrypted and cannot be processed without a password")
            
            total_pages = len(reader.pages)
            self.logger.debug(f"PDF has {total_pages} pages")
            
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    
                    self.logger.debug(
                        f"Page {page_num + 1} text extraction",
                        page_num=page_num + 1,
                        text_length=len(page_text),
                        has_text=bool(page_text.strip())
                    )
                    
                    if page_text.strip():  # Only add non-empty pages
                        text_parts.append(page_text)
                        
                except Exception as page_error:
                    self.logger.warning(
                        f"Failed to extract text from page {page_num + 1}",
                        page_num=page_num + 1,
                        error=str(page_error)
                    )
                    continue
            
            text = "\n\n".join(text_parts)
            
            self.logger.info(
                "PDF text extraction completed",
                total_pages=total_pages,
                pages_with_text=len(text_parts),
                pages_without_text=total_pages - len(text_parts),
                total_text_length=len(text),
                extraction_successful=len(text) > 0
            )
            
            # Provide helpful information if no text was extracted
            if not text.strip():
                if total_pages > 0:
                    self.logger.warning(
                        "PDF appears to be image-based or scanned",
                        total_pages=total_pages,
                        suggestion="Consider using OCR tools to extract text from image-based PDFs"
                    )
                    raise ValueError(
                        "No text content extracted from PDF. This might be an image-based/scanned PDF "
                        "that requires OCR (Optical Character Recognition) to extract text. "
                        "Try converting the PDF to a text-based format first."
                    )
                else:
                    raise ValueError("PDF has no pages")
            
            return text
            
        except Exception as e:
            self.logger.error("PDF text extraction failed", error=str(e))
            # Re-raise ValueError with PDF-specific context, wrap other exceptions
            if isinstance(e, ValueError):
                raise
            else:
                raise ValueError(f"Failed to process PDF file: {str(e)}") from e
    
    async def _extract_docx_text(self, buffer: io.BytesIO) -> str:
        """Extract text from DOCX using python-docx."""
        try:
            doc = DocxDocument(buffer)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Only add non-empty paragraphs
                    text_parts.append(paragraph.text)
            
            text = "\n".join(text_parts)
            
            self.logger.debug(
                "DOCX text extracted",
                total_paragraphs=len(doc.paragraphs),
                non_empty_paragraphs=len(text_parts),
                total_text_length=len(text)
            )
            
            # Handle DOCX-specific error case
            if not text.strip():
                total_paragraphs = len(doc.paragraphs)
                self.logger.warning(
                    "DOCX contains no extractable text",
                    total_paragraphs=total_paragraphs,
                    suggestion="Document might be empty or contain only images/graphics"
                )
                raise ValueError(
                    "No text content extracted from DOCX file. The document might be empty "
                    "or contain only images/graphics."
                )
            
            return text
            
        except Exception as e:
            self.logger.error("DOCX text extraction failed", error=str(e))
            # Re-raise ValueError with DOCX-specific context, wrap other exceptions
            if isinstance(e, ValueError):
                raise
            else:
                raise ValueError(f"Failed to process DOCX file: {str(e)}") from e
    
    async def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    self.logger.debug(
                        "TXT text extracted",
                        encoding_used=encoding,
                        text_length=len(text)
                    )
                    
                    # Handle TXT-specific error case
                    if not text.strip():
                        self.logger.warning(
                            "TXT file is empty or contains only whitespace",
                            file_size_bytes=len(file_content)
                        )
                        raise ValueError("TXT file is empty or contains no readable text content")
                    
                    return text
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error handling
            text = file_content.decode('utf-8', errors='replace')
            self.logger.warning(
                "TXT extraction used fallback encoding",
                text_length=len(text)
            )
            
            # Handle TXT-specific error case for fallback
            if not text.strip():
                raise ValueError("TXT file contains no readable text content (encoding issues detected)")
            
            return text
            
        except Exception as e:
            self.logger.error("TXT text extraction failed", error=str(e))
            # Re-raise ValueError with TXT-specific context, wrap other exceptions
            if isinstance(e, ValueError):
                raise
            else:
                raise ValueError(f"Failed to process TXT file: {str(e)}") from e
    
    async def chunk_document(self, text: str, source: str, document_id: str = None, file_metadata: Optional[Dict[str, Any]] = None) -> List[LangChainDocument]:
        """
        Split document text into LangChain Document chunks with complete metadata.
        
        Args:
            text: Document text to chunk
            source: Source document identifier
            document_id: Optional document ID for metadata
            file_metadata: Complete file metadata to include in each chunk
            
        Returns:
            List of LangChain Document objects with full metadata
        """
        try:
            if not text.strip():
                self.logger.warning("Empty text provided for chunking", source=source)
                return []
            
            # Split text using the configured text splitter
            chunks = self.text_splitter.split_text(text)
            
            # Create LangChain Document objects with complete metadata
            langchain_documents = []
            for i, chunk_content in enumerate(chunks):
                if chunk_content.strip():  # Only create chunks with content
                    metadata = {
                        "chunk_id": f"{source}_{i}",
                        "source": source,
                        "chunk_index": i,
                        "chunk_length": len(chunk_content),
                        "created_at": datetime.now(timezone.utc).isoformat()
                    }
                    
                    # Add document ID if provided
                    if document_id:
                        metadata["document_id"] = document_id
                    
                    # Add complete file metadata if provided
                    if file_metadata:
                        metadata.update(file_metadata)
                    
                    document = LangChainDocument(
                        page_content=chunk_content,
                        metadata=metadata
                    )
                    langchain_documents.append(document)
            
            self.logger.info(
                "Document chunking completed",
                source=source,
                original_length=len(text),
                total_chunks=len(langchain_documents),
                avg_chunk_size=sum(len(doc.page_content) for doc in langchain_documents) / len(langchain_documents) if langchain_documents else 0
            )
            
            return langchain_documents
            
        except Exception as e:
            self.logger.error(
                "Document chunking failed",
                source=source,
                text_length=len(text),
                error=str(e)
            )
            raise
    
    async def process_file(
        self, 
        file_path: Union[str, Path], 
        file_content: Optional[bytes] = None,
        source_name: Optional[str] = None
    ) -> List[LangChainDocument]:
        """
        Process a file completely: validate, extract text, and create LangChain Document chunks.
        All file metadata is embedded in each chunk's metadata.
        
        Args:
            file_path: Path to the file
            file_content: File content as bytes (optional)
            source_name: Custom source name (optional, uses filename if not provided)
            
        Returns:
            List of LangChain Document objects with complete metadata
        """
        file_path = Path(file_path)
        processing_start = datetime.now(timezone.utc)
        
        # Create document metadata
        document_id = str(uuid.uuid4())
        source_name = source_name or file_path.name
        
        # Prepare file metadata to embed in each chunk
        file_metadata = {
            "document_id": document_id,
            "filename": file_path.name,
            "file_type": file_path.suffix.lower(),
            "file_size": len(file_content) if file_content else file_path.stat().st_size,
            "upload_timestamp": processing_start.isoformat(),
            "processing_status": "processing"
        }
        
        try:
            # Validate file (allow memory-only processing if file_content is provided)
            self.validate_file(file_path, file_metadata["file_size"], allow_memory_only=(file_content is not None))
            
            self.logger.info(
                "Starting file processing",
                document_id=document_id,
                filename=file_path.name,
                file_size_mb=file_metadata["file_size"] / (1024*1024)
            )
            
            # Extract text - file-type specific errors are handled in respective methods
            text = await self.extract_text(file_path, file_content)
            
            # Update processing status
            file_metadata["processing_status"] = "completed"
            
            # Create LangChain Document chunks with complete metadata
            chunks = await self.chunk_document(text, source_name, document_id, file_metadata)
            
            if not chunks:
                file_metadata["processing_status"] = "failed"
                file_metadata["error_message"] = "No valid chunks created from document"
                raise ValueError("No valid chunks created from document")
            
            # Add chunk count to metadata
            for chunk in chunks:
                chunk.metadata["chunk_count"] = len(chunks)
            
            processing_time = (datetime.now(timezone.utc) - processing_start).total_seconds()
            
            self.logger.info(
                "File processing completed successfully",
                document_id=document_id,
                filename=file_path.name,
                chunk_count=len(chunks),
                processing_time_seconds=processing_time
            )
            
            return chunks
            
        except Exception as e:
            # Update status in metadata and add error info
            file_metadata["processing_status"] = "failed"
            file_metadata["error_message"] = str(e)
            
            self.logger.error(
                "File processing failed",
                document_id=document_id,
                filename=file_path.name,
                error=str(e),
                error_type=type(e).__name__
            )
            
            raise
    
    async def process_multiple_files(
        self, 
        file_data: List[tuple[Union[str, Path], Optional[bytes]]]
    ) -> List[LangChainDocument]:
        """
        Process multiple files concurrently.
        
        Args:
            file_data: List of (file_path, file_content) tuples
            
        Returns:
            List of all LangChain Document chunks from all files
        """
        self.logger.info(
            "Starting batch file processing",
            file_count=len(file_data)
        )
        
        # Process files concurrently
        tasks = [
            self.process_file(file_path, file_content)
            for file_path, file_content in file_data
        ]
        
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        all_chunks = []
        successful_count = 0
        failed_count = 0
        
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                failed_count += 1
                self.logger.error(
                    "Failed to process file in batch",
                    filename=str(file_data[i][0]),
                    error=str(result)
                )
            else:
                successful_count += 1
                all_chunks.extend(result)
        
        self.logger.info(
            "Batch file processing completed",
            successful_count=successful_count,
            failed_count=failed_count,
            total_chunks=len(all_chunks)
        )
        
